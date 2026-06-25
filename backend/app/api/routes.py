"""
FastAPI layer -- exposes the investigation orchestrator over HTTP so a frontend
(or any client) can drive it. The agents/orchestrator are unchanged; this just
translates HTTP requests into graph calls.

Endpoints
  GET  /health                              -> liveness check
  GET  /scenarios                           -> demo alerts (for the case picker)
  POST /investigate                         -> run a case (pauses at HITL)
  GET  /cases                               -> list all investigated cases
  GET  /case/{id}                           -> full current state of a case
  GET  /case/{id}/audit                     -> the audit timeline
  GET  /case/{id}/sar                       -> the SAR draft text
  POST /case/{id}/decision                  -> approve / reject / edit / request_more_info
  POST /case/{id}/rerun-agent/{agent_name}  -> re-run a single agent on the case
  POST /case/{id}/export                    -> download the case as a Markdown report

State is kept per case by LangGraph's checkpointer (thread_id = case_id), so the
pause/resume works across separate HTTP requests within the running server.
"""

import app.core.config  # noqa: F401 -- loads .env first

import json
import logging
from typing import Literal, Optional

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import Response, StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langgraph.types import Command

logger = logging.getLogger("compliguard.api")

from app.orchestrator import build_graph
from app.data.scenarios import SCENARIOS
from app.rules.rule_engine import get_rules, reload_rules
from app.rules.country_risk import get_country_risk
from app.tools.rag import load_policies, reset_policy_collection
from app.services.persistence import (
    persist_case, persist_decision, list_cases, get_case_events, get_case_sar,
)

# agent instances that are safe to re-run standalone (no human interrupt)
from app.agents.alert_intake import alert_intake
from app.agents.data_quality import data_quality
from app.agents.transaction_analysis import transaction_analysis
from app.agents.kyc_profile import kyc_profile
from app.agents.watchlist_screening import watchlist_screening
from app.agents.policy_rag import policy_rag
from app.agents.case_memory import case_memory
from app.agents.risk_scoring import risk_scoring

AGENTS = {a.name: a for a in [
    alert_intake, data_quality, transaction_analysis, kyc_profile,
    watchlist_screening, policy_rag, case_memory, risk_scoring,
]}

# human-readable labels for streamed progress events
LABELS = {
    "alert_intake": "Alert Intake Agent",
    "data_quality": "Data Quality Gate",
    "transaction_analysis": "Transaction Analysis Agent",
    "kyc_profile": "KYC Profile Agent",
    "watchlist_screening": "Watchlist Screening Agent",
    "policy_rag": "Policy RAG Agent",
    "case_memory": "Memory Agent",
    "risk_scoring": "Risk Scoring Agent",
    "sar_drafting": "SAR Drafting Agent",
    "compliance_review": "Compliance Review Agent",
    "human_approval": "Human Approval",
}

app = FastAPI(title="CompliGuard AI", version="1.0",
              description="Multi-agent AML compliance investigation API")

app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"],
)


@app.exception_handler(Exception)
async def _unhandled(request: Request, exc: Exception):
    """Any unexpected error returns a clean JSON 500 (no leaked traceback)."""
    logger.exception(f"Unhandled error on {request.method} {request.url.path}")
    return JSONResponse(status_code=500,
                        content={"detail": f"Internal error: {type(exc).__name__}"})


# one graph for the process; the checkpointer holds each case's state in memory
graph = build_graph()


# ======================================================================
# Request / response models  (typed -> self-documenting in /docs)
# ======================================================================
class Alert(BaseModel):
    id: str
    customer_id: str
    reason: str
    recipient: Optional[str] = ""
    country: Optional[str] = ""
    total_amount: Optional[int] = 0
    num_transactions: Optional[int] = 0


class Decision(BaseModel):
    decision: Literal["approve", "reject", "request_more_info", "edit"]
    analyst_id: str
    analyst_note: Optional[str] = None
    edited_sar_draft: Optional[str] = None
    final_risk_level: Optional[str] = None
    rerun_targets: Optional[list[str]] = None


class CaseSnapshot(BaseModel):
    case_id: str
    status: str
    triage: Optional[dict] = None
    data_quality: Optional[dict] = None
    transaction_findings: Optional[dict] = None
    kyc_findings: Optional[dict] = None
    watchlist_findings: Optional[dict] = None
    memory_findings: Optional[dict] = None
    retrieved_policies: Optional[list] = None
    risk_score: Optional[int] = None
    rule_score: Optional[int] = None
    ai_score: Optional[int] = None
    risk_level: Optional[str] = None
    risk_factors: Optional[list] = None
    key_drivers: Optional[list] = None
    recommendation: Optional[str] = None
    risk_explanation: Optional[str] = None
    sar_draft: Optional[str] = None
    review: Optional[dict] = None
    human_decision: Optional[str] = None
    human_review: Optional[dict] = None
    errors: Optional[list] = None
    audit: Optional[list] = None
    cot_traces: Optional[list] = None
    a2a_messages: Optional[list] = None


class CaseSummary(BaseModel):
    case_id: str
    customer_id: Optional[str] = None
    typology: Optional[str] = None
    status: Optional[str] = None
    risk_score: Optional[int] = None
    risk_level: Optional[str] = None
    updated_at: Optional[str] = None


class HealthResponse(BaseModel):
    status: str


class AuditResponse(BaseModel):
    case_id: str
    timeline: list[str]
    events: list[dict]


class SARResponse(BaseModel):
    case_id: str
    exists: bool
    sar_draft: Optional[str] = None


# ======================================================================
# Helpers
# ======================================================================
def _state(case_id: str) -> dict:
    v = graph.get_state({"configurable": {"thread_id": case_id}}).values
    if not v:
        raise HTTPException(status_code=404, detail=f"No case '{case_id}'")
    return v


def _snapshot(case_id: str) -> dict:
    """Serialize a case's current state."""
    cfg = {"configurable": {"thread_id": case_id}}
    state = graph.get_state(cfg)
    v = state.values
    if not v:
        raise HTTPException(status_code=404, detail=f"No case '{case_id}'")

    dq = v.get("data_quality", {})
    awaiting = "human_approval" in (state.next or ())
    if dq and not dq.get("complete", True):
        status = "needs_more_information"
    elif v.get("errors") and not v.get("human_decision"):
        status = "manual_review_required"
    elif awaiting:
        status = "awaiting_decision"
    elif v.get("human_decision"):
        status = "closed"
    else:
        status = "auto_closed"

    return {
        "case_id": case_id, "status": status,
        "triage": v.get("triage"), "data_quality": v.get("data_quality"),
        "transaction_findings": v.get("transaction_findings"),
        "kyc_findings": v.get("kyc_findings"),
        "watchlist_findings": v.get("watchlist_findings"),
        "memory_findings": v.get("memory_findings"),
        "retrieved_policies": v.get("retrieved_policies"),
        "risk_score": v.get("risk_score"), "rule_score": v.get("rule_score"),
        "ai_score": v.get("ai_score"), "risk_level": v.get("risk_level"),
        "risk_factors": v.get("risk_factors"), "key_drivers": v.get("key_drivers"),
        "recommendation": v.get("recommendation"),
        "risk_explanation": v.get("risk_explanation"),
        "sar_draft": v.get("sar_draft"), "review": v.get("review"),
        "human_decision": v.get("human_decision"),
        "human_review": v.get("human_review"), "errors": v.get("errors"),
        "audit": sorted(v.get("audit", [])),
        "cot_traces": v.get("cot_traces"), "a2a_messages": v.get("a2a_messages"),
    }


def _report_sections(snap: dict) -> list[tuple[str, list[str]]]:
    """The full SAR report content as (heading, lines) sections."""
    tri = snap.get("triage") or {}
    tf = snap.get("transaction_findings") or {}
    kyc = snap.get("kyc_findings") or {}
    wl = snap.get("watchlist_findings") or {}
    mem = snap.get("memory_findings") or {}
    hr = snap.get("human_review") or {}
    cust_id = (tri.get("entities") or {}).get("customer_id", "-")

    sections = [
        ("Case Overview", [
            f"Case ID: {snap.get('case_id')}",
            f"Customer ID: {cust_id}",
            f"Status: {snap.get('status')}",
            f"Alert type: {tri.get('alert_type')}  |  Priority: {tri.get('priority')}",
        ]),
        ("Risk Assessment", [
            f"Final score: {snap.get('risk_score')}/100 ({snap.get('risk_level')})",
            f"Rule baseline: {snap.get('rule_score')}  |  AI assessment: {snap.get('ai_score')}",
            f"Recommendation: {snap.get('recommendation')}",
            f"Explanation: {snap.get('risk_explanation')}",
        ]),
    ]
    if snap.get("risk_factors"):
        sections.append(("Triggered AML Rules", [
            f"[{f.get('rule_id', '')}] {f.get('points'):+} {f.get('name')} "
            f"({f.get('severity', '')}) - {f.get('evidence')}"
            for f in snap["risk_factors"]]))
    sections.append(("Agent Findings", [
        f"Typology: {tf.get('typology')}",
        f"KYC: {kyc.get('consistency')} | checks failed: {kyc.get('checks_failed')} "
        f"| EDD: {kyc.get('edd_required')}",
        f"Watchlist: {wl.get('verdict')} (best {wl.get('match_score')}% on {wl.get('list_type')})",
        f"Memory: {mem.get('memory_risk_signal')}",
    ]))
    if snap.get("retrieved_policies"):
        sections.append(("Policy References", [
            f"{p.get('policy_id')}: {p.get('title')} (section {p.get('section')})"
            for p in snap["retrieved_policies"]]))
    if snap.get("sar_draft"):
        sections.append(("SAR Draft", snap["sar_draft"].split("\n")))
    if hr:
        sections.append(("Human Decision", [
            f"Decision: {hr.get('decision')}",
            f"Analyst: {hr.get('analyst_id')}",
            f"Note: {hr.get('analyst_note') or '-'}",
            f"Risk-level override: {hr.get('final_risk_level') or '-'}",
        ]))
    elif snap.get("human_decision"):
        sections.append(("Human Decision", [f"Decision: {snap.get('human_decision')}"]))
    if snap.get("audit"):
        sections.append(("Audit Timeline", list(snap["audit"])))
    return sections


def _report_md(snap: dict) -> str:
    out = [f"# Suspicious Activity Report - {snap.get('case_id')}", ""]
    for title, lines in _report_sections(snap):
        out.append(f"## {title}")
        out.append("")
        out += (lines if title == "SAR Draft" else [f"- {l}" for l in lines])
        out.append("")
    return "\n".join(out)


def _ascii(s) -> str:
    """fpdf core fonts are Latin-1 only -- replace common unicode."""
    return (str(s).replace("—", "-").replace("–", "-").replace("→", "->")
            .replace("✓", "[OK]").replace("…", "...").replace("’", "'")
            .encode("latin-1", "replace").decode("latin-1"))


def _report_pdf(snap: dict) -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def line(text, size=10, bold=False, gap=0):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        pdf.set_x(pdf.l_margin)
        if gap:
            pdf.ln(gap)
        pdf.multi_cell(pdf.epw, size * 0.55 + 1, _ascii(text) or " ", wrapmode="CHAR")

    line(f"Suspicious Activity Report - {snap.get('case_id')}", size=16, bold=True)
    for title, lines in _report_sections(snap):
        line(title, size=13, bold=True, gap=3)
        for l in lines:
            line(l, size=10)
    return bytes(pdf.output())


def _report_docx(snap: dict) -> bytes:
    from io import BytesIO
    from docx import Document
    doc = Document()
    doc.add_heading(f"Suspicious Activity Report - {snap.get('case_id')}", level=0)
    for title, lines in _report_sections(snap):
        doc.add_heading(title, level=1)
        for l in lines:
            doc.add_paragraph(str(l))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_report(snap: dict) -> str:
    """Backwards-compatible Markdown report (used by POST /case/{id}/export)."""
    return _report_md(snap)


# ======================================================================
# Endpoints
# ======================================================================
@app.get("/health", response_model=HealthResponse)
def health():
    return {"status": "ok"}


@app.get("/health/ready")
def ready():
    """Deep readiness check -- confirms the LLM (Ollama) and database (Supabase)
    are actually reachable. Useful before a live demo."""
    deps = {}
    try:
        from app.services.llm import _client
        _client.models.list()
        deps["ollama"] = "ok"
    except Exception as exc:
        deps["ollama"] = f"error: {type(exc).__name__}"
    try:
        from app.tools.db import client
        client().table("customers").select("customer_id").limit(1).execute()
        deps["supabase"] = "ok"
    except Exception as exc:
        deps["supabase"] = f"error: {type(exc).__name__}"
    return {"ready": all(v == "ok" for v in deps.values()), "dependencies": deps}


@app.get("/scenarios")
def scenarios():
    return SCENARIOS


@app.get("/policies")
def policies():
    """List the policy documents currently available to the RAG layer."""
    return [{"policy_id": p["id"], "title": p["title"], "section": p["section"],
             "category": p["category"]} for p in load_policies()]


@app.post("/policies/reindex")
def policies_reindex():
    """Re-index the policy folder -- picks up newly added or edited .md/.pdf
    policy files WITHOUT restarting the server."""
    reset_policy_collection()
    docs = load_policies()
    return {"reindexed": len(docs), "policies": [p["id"] for p in docs]}


@app.get("/rules")
def rules():
    """The institution's configurable AML rule set (thresholds + risk points)."""
    return get_rules()


@app.get("/country-risk")
def country_risk():
    """The country-risk register (level, reason, source, last reviewed)."""
    return get_country_risk()


@app.get("/watchlist")
def watchlist():
    """The active watchlist entities across all lists (sanctions, PEP, blacklist,
    adverse media, scam/mule accounts, high-risk entities)."""
    from app.tools.db import get_watchlist
    return get_watchlist()


@app.post("/rules/reload")
def rules_reload():
    """Re-read risk_rules.yaml after editing thresholds -- no restart needed."""
    return reload_rules()


@app.post("/investigate", response_model=CaseSnapshot)
def investigate(alert: Alert):
    """Run the full investigation. Returns when the graph pauses for human
    approval (high risk) or ends (low risk / incomplete / tool failure)."""
    case_id = alert.id
    cfg = {"configurable": {"thread_id": case_id}}
    graph.invoke({"alert": alert.model_dump()}, cfg)
    snap = _snapshot(case_id)
    persist_case(graph.get_state(cfg).values, status=snap["status"])
    return snap


@app.post("/investigate/stream")
def investigate_stream(alert: Alert):
    """Run the investigation and STREAM agent progress as Server-Sent Events.
    Each agent emits a 'progress' event as it completes; a final 'done' event
    carries the status. Consume with fetch()+ReadableStream on the frontend.

    Event stream:
        event: progress   data: {agent, label, status, confidence, message}
        ...
        event: done        data: {case_id, status, risk_score, risk_level}
    """
    case_id = alert.id
    cfg = {"configurable": {"thread_id": case_id}}

    def gen():
        try:
            for chunk in graph.stream({"alert": alert.model_dump()}, cfg,
                                      stream_mode="updates"):
                for node, updates in chunk.items():
                    if node == "__interrupt__" or not isinstance(updates, dict):
                        continue
                    conf = None
                    if updates.get("cot_traces"):
                        conf = updates["cot_traces"][-1].get("confidence")
                    msg = updates["audit"][-1] if updates.get("audit") else ""
                    ev = {"agent": node, "label": LABELS.get(node, node),
                          "status": "completed", "confidence": conf, "message": msg}
                    yield f"event: progress\ndata: {json.dumps(ev)}\n\n"

            # investigation finished (paused for human, or ended)
            snap = _snapshot(case_id)
            persist_case(graph.get_state(cfg).values, status=snap["status"])
            done = {"case_id": case_id, "status": snap["status"],
                    "risk_score": snap.get("risk_score"),
                    "risk_level": snap.get("risk_level")}
            yield f"event: done\ndata: {json.dumps(done)}\n\n"
        except Exception as exc:
            yield f"event: error\ndata: {json.dumps({'error': str(exc)})}\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream")


@app.get("/cases", response_model=list[CaseSummary])
def cases(limit: int = 50):
    """List investigated cases (survives server restarts)."""
    return list_cases(limit)


@app.get("/case/{case_id}/status")
def case_status(case_id: str):
    """Lightweight status poll (for clients not using the SSE stream)."""
    snap = _snapshot(case_id)
    return {"case_id": case_id, "status": snap["status"],
            "risk_score": snap.get("risk_score"), "risk_level": snap.get("risk_level")}


@app.get("/case/{case_id}", response_model=CaseSnapshot)
def get_case(case_id: str):
    return _snapshot(case_id)


@app.get("/case/{case_id}/audit", response_model=AuditResponse)
def case_audit(case_id: str):
    """The audit timeline -- live (in memory) plus persisted events."""
    v = graph.get_state({"configurable": {"thread_id": case_id}}).values
    timeline = sorted(v.get("audit", [])) if v else []
    return {"case_id": case_id, "timeline": timeline, "events": get_case_events(case_id)}


@app.get("/case/{case_id}/sar", response_model=SARResponse)
def case_sar(case_id: str):
    """The SAR draft text (live state, falling back to the persisted copy)."""
    v = graph.get_state({"configurable": {"thread_id": case_id}}).values
    draft = (v or {}).get("sar_draft") or get_case_sar(case_id)
    return {"case_id": case_id, "exists": bool(draft), "sar_draft": draft}


@app.post("/case/{case_id}/decision", response_model=CaseSnapshot)
def decide(case_id: str, body: Decision):
    """Resume a paused case with the analyst's structured decision.
    'request_more_info' re-runs the investigation and pauses again."""
    cfg = {"configurable": {"thread_id": case_id}}
    graph.invoke(Command(resume=body.model_dump()), cfg)
    persist_decision(case_id, body.decision, analyst_id=body.analyst_id,
                     notes=body.analyst_note, final_risk_level=body.final_risk_level)
    snap = _snapshot(case_id)
    persist_case(graph.get_state(cfg).values, status=snap["status"])
    return snap


@app.post("/case/{case_id}/rerun-agent/{agent_name}", response_model=CaseSnapshot)
def rerun_agent(case_id: str, agent_name: str):
    """Re-run a single agent (e.g. watchlist_screening, policy_rag) on an existing
    case and merge its fresh output back in -- handy when a tool was flaky."""
    if agent_name not in AGENTS:
        raise HTTPException(status_code=404,
                            detail=f"Unknown agent '{agent_name}'. "
                                   f"Options: {sorted(AGENTS)}")
    cfg = {"configurable": {"thread_id": case_id}}
    v = _state(case_id)
    updates = AGENTS[agent_name](v)         # run just that agent
    graph.update_state(cfg, updates)        # merge into the checkpoint
    persist_case(graph.get_state(cfg).values, status=_snapshot(case_id)["status"])
    return _snapshot(case_id)


@app.post("/case/{case_id}/export")
def export_case(case_id: str):
    """Download the case as a Markdown report."""
    report = _build_report(_snapshot(case_id))
    return Response(
        content=report, media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{case_id}_report.md"'},
    )


@app.post("/case/{case_id}/export-sar")
def export_sar(case_id: str, format: Literal["pdf", "markdown", "docx"] = "pdf"):
    """Download the full SAR report as PDF, Markdown, or DOCX. Includes case info,
    risk score + factors, agent findings, policy references, the SAR draft, the
    human decision, and the audit timeline."""
    snap = _snapshot(case_id)
    builders = {
        "pdf":      (_report_pdf,  "application/pdf", "pdf"),
        "docx":     (_report_docx, "application/vnd.openxmlformats-officedocument."
                                   "wordprocessingml.document", "docx"),
        "markdown": (lambda s: _report_md(s).encode(), "text/markdown", "md"),
    }
    build, media, ext = builders[format]
    return Response(
        content=build(snap), media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{case_id}_SAR.{ext}"'},
    )
