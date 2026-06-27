"""
Document text extraction for the 'upload an alert / statement' flow.

Pulls plain text out of an uploaded PDF / DOCX / TXT / CSV so the LLM can extract
the alert, the customer profile, and the transactions.

For PDFs we don't use a naive get_text() -- on a real bank statement that flattens
the table into a single column and the debit/credit distinction is lost. Instead we
reconstruct the visual ROWS from word coordinates (cluster by y, order by x), so each
transaction becomes one line: `date  description  amount  running-balance`. The LLM can
then derive direction from the running-balance delta, which works across bank layouts.

Best-effort: unsupported or image-only (scanned) files return an empty string, and the
caller surfaces a clear "scanned PDF / no text layer" message.
"""

import logging
from io import BytesIO

logger = logging.getLogger(__name__)

_Y_BAND = 3.0     # points; lines are ~17pt apart, so this clusters a row without merging rows
_OCR_Y_BAND = 14  # pixels at ~200 dpi; row tolerance for OCR'd boxes
_OCR_MAX_PAGES = 12
_ocr_reader = None  # lazily-initialised EasyOCR reader (expensive to construct)


def _rows_from_bands(bands: dict) -> list[str]:
    """Join position-keyed tokens into one line per visual row (y-band)."""
    out = []
    for key in sorted(bands):
        line = " ".join(w for _, w in sorted(bands[key], key=lambda p: p[0])).strip()
        if line:
            out.append(line)
    return out


def _pdf_rows(data: bytes) -> str:
    """Reconstruct visual rows from word positions, page by page (digital PDFs)."""
    import fitz  # PyMuPDF

    out: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            words = page.get_text("words")  # (x0, y0, x1, y1, text, block, line, word_no)
            if not words:
                continue
            bands: dict[int, list[tuple[float, str]]] = {}
            for x0, y0, _x1, _y1, text, *_ in words:
                if text.strip():
                    bands.setdefault(round(y0 / _Y_BAND), []).append((x0, text))
            out.extend(_rows_from_bands(bands))
    return "\n".join(out)


def _get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr  # heavy import; only when a scanned PDF is actually uploaded
        import torch
        # verbose=False suppresses the block-char progress bar that crashes a cp1252 console
        _ocr_reader = easyocr.Reader(["en"], gpu=torch.cuda.is_available(), verbose=False)
    return _ocr_reader


def ocr_available() -> bool:
    import importlib.util
    return importlib.util.find_spec("easyocr") is not None


def _ocr_pdf_rows(data: bytes) -> str:
    """OCR an image-only (scanned) PDF, reconstructing rows from word boxes so the
    same statement parser works. Best-effort: returns '' if OCR isn't available."""
    if not ocr_available():
        return ""
    try:
        import fitz
        reader = _get_ocr_reader()
        out: list[str] = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in list(doc)[:_OCR_MAX_PAGES]:
                pix = page.get_pixmap(dpi=200)
                results = reader.readtext(pix.tobytes("png"))  # [(bbox, text, conf), ...]
                bands: dict[int, list[tuple[float, str]]] = {}
                for bbox, text, conf in results:
                    if not text.strip() or conf < 0.3:
                        continue
                    ys = [pt[1] for pt in bbox]
                    xs = [pt[0] for pt in bbox]
                    bands.setdefault(round((sum(ys) / len(ys)) / _OCR_Y_BAND), []).append(
                        (min(xs), text.strip()))
                out.extend(_rows_from_bands(bands))
        return "\n".join(out)
    except Exception as exc:
        logger.warning(f"[doc_extract] OCR failed: {exc}")
        return ""


def extract_text(filename: str, data: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in (filename or "") else ""
    try:
        if ext == "pdf":
            text = _pdf_rows(data).strip()
            if text:
                return text
            # no text layer -> scanned / image-only PDF: fall back to OCR
            logger.info("[doc_extract] no text layer; attempting OCR")
            return _ocr_pdf_rows(data).strip()
        if ext == "docx":
            from docx import Document
            d = Document(BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.append("  ".join(c.text for c in row.cells))
            return "\n".join(parts).strip()
        if ext in ("txt", "md", "csv", "eml", "json"):
            return data.decode("utf-8", "ignore").strip()
        # unknown extension -> try to decode as text
        return data.decode("utf-8", "ignore").strip()
    except Exception as exc:
        logger.warning(f"[doc_extract] could not read {filename}: {exc}")
        return ""
