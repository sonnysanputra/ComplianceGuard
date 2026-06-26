"""
Professional SAR report rendering (PDF + DOCX).

Produces a branded, regulator-style document: a header band with the
ComplianceGuard logo, a CONFIDENTIAL classification, a metadata block, cleanly
styled numbered sections, and a footer with page numbers + a draft disclaimer.

The logo is used if app/assets/logo.png exists; otherwise a drawn shield + the
"ComplianceGuard" wordmark is rendered instead, so output is professional either way.
"""

from io import BytesIO
from pathlib import Path

ASSETS = Path(__file__).resolve().parent.parent / "assets"


def _logo() -> Path | None:
    """The brand logo, looked up by either expected filename (or None)."""
    for name in ("logo.png", "ComplianceGuardLogo.png", "ComplianceGuard.png"):
        p = ASSETS / name
        if p.exists():
            return p
    return None

# brand palette
ORANGE = (249, 115, 22)
DARK = (31, 41, 55)
GREY = (107, 114, 128)
LIGHT = (229, 231, 235)
SOFT = (248, 250, 252)
WHITE = (255, 255, 255)

DISCLAIMER = ("AI-generated DRAFT investigation package. Requires human analyst "
              "review and sign-off before any STR is lodged with FIED.")


def _ascii(s) -> str:
    """fpdf core fonts are Latin-1 only -- replace common unicode."""
    return (str(s).replace("—", "-").replace("–", "-").replace("→", "->")
            .replace("•", "-").replace("✓", "[OK]").replace("…", "...")
            .replace("’", "'").replace("“", '"').replace("”", '"')
            .encode("latin-1", "replace").decode("latin-1"))


# ======================================================================
# PDF
# ======================================================================
def render_pdf(title: str, subtitle: str, meta: list[tuple[str, str]],
               sections: list[tuple[str, list[str]]]) -> bytes:
    from fpdf import FPDF
    logo = _logo()

    class Doc(FPDF):
        def header(self):
            # brand band
            if logo:
                try:
                    self.image(str(logo), x=15, y=9, h=10)
                    bx = 27
                except Exception:
                    bx = 15
            else:
                self.set_xy(15, 9)
                self.set_font("Helvetica", "B", 13)
                self.set_text_color(*DARK); self.cell(self.get_string_width("Compliance"), 8, "Compliance")
                self.set_text_color(*ORANGE); self.cell(20, 8, "Guard")
                bx = None
            if bx is not None:
                self.set_xy(bx, 10)
                self.set_font("Helvetica", "B", 12)
                self.set_text_color(*DARK); self.cell(self.get_string_width("Compliance"), 8, "Compliance")
                self.set_text_color(*ORANGE); self.cell(20, 8, "Guard")
            # classification tag (right)
            self.set_font("Helvetica", "B", 7.5)
            self.set_text_color(*ORANGE)
            self.set_xy(-60, 11)
            self.cell(45, 6, "CONFIDENTIAL", align="R")
            # rule line
            self.set_draw_color(*ORANGE); self.set_line_width(0.6)
            self.line(15, 22, self.w - 15, 22)
            self.set_y(28)

        def footer(self):
            self.set_y(-14)
            self.set_draw_color(*LIGHT); self.set_line_width(0.2)
            self.line(15, self.get_y(), self.w - 15, self.get_y())
            self.set_font("Helvetica", "", 7.5); self.set_text_color(*GREY)
            self.set_y(-12)
            self.cell(0, 6, "ComplianceGuard AI  -  Confidential: for internal compliance review only", align="L")
            self.set_y(-12)
            self.cell(0, 6, f"Page {self.page_no()} of {{nb}}", align="R")

    pdf = Doc()
    pdf.alias_nb_pages()
    pdf.set_margins(15, 28, 15)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # ---- title block ----
    pdf.set_font("Helvetica", "B", 20); pdf.set_text_color(*DARK)
    pdf.set_x(15); pdf.multi_cell(pdf.epw, 9, _ascii(title))
    pdf.set_font("Helvetica", "", 10.5); pdf.set_text_color(*GREY)
    pdf.set_x(15); pdf.multi_cell(pdf.epw, 5.5, _ascii(subtitle))
    pdf.ln(2)

    # ---- metadata table ----
    pdf.set_font("Helvetica", "", 9)
    col_l, col_r = 38, pdf.epw - 38
    for i, (k, v) in enumerate(meta):
        pdf.set_x(15)
        pdf.set_fill_color(*(SOFT if i % 2 == 0 else WHITE))
        pdf.set_text_color(*GREY); pdf.set_font("Helvetica", "B", 8.5)
        pdf.cell(col_l, 7, "  " + _ascii(k), border=0, fill=True)
        pdf.set_text_color(*DARK); pdf.set_font("Helvetica", "", 9)
        pdf.cell(col_r, 7, _ascii(v), border=0, fill=True, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # ---- sections ----
    for title_s, lines in sections:
        if pdf.get_y() > pdf.h - 40:
            pdf.add_page()
        pdf.set_x(15)
        pdf.set_font("Helvetica", "B", 11.5); pdf.set_text_color(*DARK)
        pdf.multi_cell(pdf.epw, 7, _ascii(title_s))
        pdf.set_draw_color(*ORANGE); pdf.set_line_width(0.4)
        y = pdf.get_y() + 0.5
        pdf.line(15, y, 15 + 24, y)
        pdf.ln(2.5)
        pdf.set_font("Helvetica", "", 9.5); pdf.set_text_color(60, 60, 60)
        for ln in lines:
            pdf.set_x(17)
            pdf.multi_cell(pdf.epw - 2, 5.2, _ascii(ln) or " ", wrapmode="CHAR")
        pdf.ln(3)

    return bytes(pdf.output())


# ======================================================================
# DOCX
# ======================================================================
def render_docx(title: str, subtitle: str, meta: list[tuple[str, str]],
                sections: list[tuple[str, list[str]]]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    O = RGBColor(0xF9, 0x73, 0x16)
    D = RGBColor(0x1F, 0x29, 0x37)
    G = RGBColor(0x6B, 0x72, 0x80)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # header band: logo + wordmark + classification
    logo = _logo()
    hp = doc.add_paragraph()
    if logo:
        try:
            hp.add_run().add_picture(str(logo), height=Inches(0.42))
            hp.add_run("   ")
        except Exception:
            pass
    r = hp.add_run("Compliance"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = D
    r = hp.add_run("Guard"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = O
    tag = hp.add_run("        CONFIDENTIAL"); tag.bold = True; tag.font.size = Pt(8); tag.font.color.rgb = O

    # title
    t = doc.add_paragraph(); tr = t.add_run(title); tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = D
    s = doc.add_paragraph(); sr = s.add_run(subtitle); sr.italic = True; sr.font.size = Pt(10.5); sr.font.color.rgb = G

    # metadata table
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for k, v in meta:
        cells = table.add_row().cells
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(9); kr.font.color.rgb = G
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(9.5); vr.font.color.rgb = D
    doc.add_paragraph()

    # sections
    for title_s, lines in sections:
        h = doc.add_heading(level=1)
        hr = h.add_run(title_s); hr.font.color.rgb = O
        for ln in lines:
            p = doc.add_paragraph(str(ln), style="List Bullet")
            p.runs[0].font.size = Pt(9.5)

    # disclaimer footer text
    doc.add_paragraph()
    fp = doc.add_paragraph(); fr = fp.add_run(DISCLAIMER); fr.italic = True; fr.font.size = Pt(8.5); fr.font.color.rgb = G

    # page-number footer
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("ComplianceGuard AI  ·  Confidential  ·  Page ")
    fr.font.size = Pt(8); fr.font.color.rgb = G
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()
